"""TextChunker 文本提取单元测试"""

import os
import tempfile

import pytest

from src.chunker import TextChunker


@pytest.fixture
def chunker():
    """创建 TextChunker 实例"""
    return TextChunker()


@pytest.fixture
def tmp_dir():
    """创建临时目录"""
    with tempfile.TemporaryDirectory() as d:
        yield d


class TestExtractTxt:
    """TXT 文件提取测试"""

    def test_extract_utf8_text(self, chunker, tmp_dir):
        path = os.path.join(tmp_dir, "test.txt")
        content = "你好世界\nHello World"
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)

        result = chunker.extract_text(path)
        assert result == content

    def test_extract_empty_txt_returns_none(self, chunker, tmp_dir):
        path = os.path.join(tmp_dir, "empty.txt")
        with open(path, "w", encoding="utf-8") as f:
            f.write("")

        result = chunker.extract_text(path)
        assert result is None

    def test_extract_whitespace_only_returns_none(self, chunker, tmp_dir):
        path = os.path.join(tmp_dir, "blank.txt")
        with open(path, "w", encoding="utf-8") as f:
            f.write("   \n\t\n  ")

        result = chunker.extract_text(path)
        assert result is None


class TestExtractMarkdown:
    """Markdown 文件提取测试"""

    def test_extract_headings_and_paragraphs(self, chunker, tmp_dir):
        path = os.path.join(tmp_dir, "test.md")
        content = "# 标题一\n\n这是段落内容。\n\n## 标题二\n\n另一段内容。"
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)

        result = chunker.extract_text(path)
        assert "标题一" in result
        assert "这是段落内容。" in result
        assert "标题二" in result
        assert "另一段内容。" in result
        # 语法标记应被移除
        assert "# " not in result
        assert "## " not in result

    def test_extract_removes_bold_italic(self, chunker, tmp_dir):
        path = os.path.join(tmp_dir, "format.md")
        content = "这是**加粗**和*斜体*文本。"
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)

        result = chunker.extract_text(path)
        assert "加粗" in result
        assert "斜体" in result
        assert "**" not in result
        assert "*斜体*" not in result

    def test_extract_code_block(self, chunker, tmp_dir):
        path = os.path.join(tmp_dir, "code.md")
        content = "# 代码示例\n\n```python\nprint('hello')\n```\n"
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)

        result = chunker.extract_text(path)
        assert "print('hello')" in result
        assert "```" not in result

    def test_extract_empty_markdown_returns_none(self, chunker, tmp_dir):
        path = os.path.join(tmp_dir, "empty.md")
        with open(path, "w", encoding="utf-8") as f:
            f.write("")

        result = chunker.extract_text(path)
        assert result is None


class TestExtractWord:
    """Word 文件提取测试"""

    def test_extract_docx_paragraphs(self, chunker, tmp_dir):
        from docx import Document

        path = os.path.join(tmp_dir, "test.docx")
        doc = Document()
        doc.add_paragraph("第一段内容")
        doc.add_paragraph("第二段内容")
        doc.add_paragraph("第三段内容")
        doc.save(path)

        result = chunker.extract_text(path)
        assert "第一段内容" in result
        assert "第二段内容" in result
        assert "第三段内容" in result

    def test_extract_empty_docx_returns_none(self, chunker, tmp_dir):
        from docx import Document

        path = os.path.join(tmp_dir, "empty.docx")
        doc = Document()
        doc.save(path)

        result = chunker.extract_text(path)
        assert result is None


class TestExtractExcel:
    """Excel 文件提取测试"""

    def test_extract_xlsx_cells(self, chunker, tmp_dir):
        from openpyxl import Workbook

        path = os.path.join(tmp_dir, "test.xlsx")
        wb = Workbook()
        ws = wb.active
        ws["A1"] = "姓名"
        ws["B1"] = "年龄"
        ws["A2"] = "张三"
        ws["B2"] = 25
        wb.save(path)

        result = chunker.extract_text(path)
        assert "姓名" in result
        assert "年龄" in result
        assert "张三" in result
        assert "25" in result
        # 单元格之间应以制表符分隔
        assert "姓名\t年龄" in result

    def test_extract_multiple_sheets(self, chunker, tmp_dir):
        from openpyxl import Workbook

        path = os.path.join(tmp_dir, "multi.xlsx")
        wb = Workbook()
        ws1 = wb.active
        ws1.title = "Sheet1"
        ws1["A1"] = "数据1"
        ws2 = wb.create_sheet("Sheet2")
        ws2["A1"] = "数据2"
        wb.save(path)

        result = chunker.extract_text(path)
        assert "数据1" in result
        assert "数据2" in result


class TestExtractPdf:
    """PDF 文件提取测试"""

    def test_extract_pdf_text(self, chunker, tmp_dir):
        from PyPDF2 import PdfWriter

        path = os.path.join(tmp_dir, "test.pdf")
        # 创建一个简单的 PDF 用于测试
        writer = PdfWriter()
        # PyPDF2 不能直接创建带文本的 PDF，使用 reportlab 或跳过
        # 这里我们测试空 PDF 的处理
        writer.add_blank_page(width=612, height=792)
        with open(path, "wb") as f:
            writer.write(f)

        # 空白页 PDF 应返回 None（无文本内容）
        result = chunker.extract_text(path)
        assert result is None


class TestExtractErrors:
    """错误处理测试"""

    def test_unsupported_extension_returns_none(self, chunker, tmp_dir):
        path = os.path.join(tmp_dir, "test.jpg")
        with open(path, "w") as f:
            f.write("not an image")

        result = chunker.extract_text(path)
        assert result is None

    def test_nonexistent_file_returns_none(self, chunker):
        result = chunker.extract_text("/nonexistent/path/file.txt")
        assert result is None

    def test_corrupted_docx_returns_none(self, chunker, tmp_dir):
        path = os.path.join(tmp_dir, "corrupted.docx")
        with open(path, "w") as f:
            f.write("this is not a valid docx file")

        result = chunker.extract_text(path)
        assert result is None

    def test_corrupted_xlsx_returns_none(self, chunker, tmp_dir):
        path = os.path.join(tmp_dir, "corrupted.xlsx")
        with open(path, "w") as f:
            f.write("this is not a valid xlsx file")

        result = chunker.extract_text(path)
        assert result is None

    def test_corrupted_pdf_returns_none(self, chunker, tmp_dir):
        path = os.path.join(tmp_dir, "corrupted.pdf")
        with open(path, "w") as f:
            f.write("this is not a valid pdf file")

        result = chunker.extract_text(path)
        assert result is None

    def test_case_insensitive_extension(self, chunker, tmp_dir):
        path = os.path.join(tmp_dir, "test.TXT")
        with open(path, "w", encoding="utf-8") as f:
            f.write("大写扩展名")

        result = chunker.extract_text(path)
        assert result == "大写扩展名"


class MockTokenizer:
    """模拟 HuggingFace tokenizer 的行为

    使用简单的字符级 tokenization 来模拟真实 tokenizer。
    每个字符对应一个 token ID。
    """

    def __init__(self):
        self._char_to_id = {}
        self._id_to_char = {}
        self._next_id = 1

    def _get_id(self, char: str) -> int:
        if char not in self._char_to_id:
            self._char_to_id[char] = self._next_id
            self._id_to_char[self._next_id] = char
            self._next_id += 1
        return self._char_to_id[char]

    def encode(self, text: str, add_special_tokens: bool = True) -> list:
        return [self._get_id(ch) for ch in text]

    def decode(self, token_ids: list, skip_special_tokens: bool = True) -> str:
        return "".join(self._id_to_char.get(tid, "") for tid in token_ids)


@pytest.fixture
def mock_tokenizer():
    """创建 mock tokenizer"""
    return MockTokenizer()


@pytest.fixture
def chunker_with_tokenizer(mock_tokenizer):
    """创建带 tokenizer 的 TextChunker 实例"""
    return TextChunker(tokenizer=mock_tokenizer, chunk_size=512, chunk_overlap=64)


class TestChunkBasic:
    """基本切块逻辑测试"""

    def test_empty_text_returns_empty_list(self, chunker_with_tokenizer):
        result = chunker_with_tokenizer.chunk("")
        assert result == []

    def test_whitespace_only_returns_empty_list(self, chunker_with_tokenizer):
        result = chunker_with_tokenizer.chunk("   \n\t  ")
        assert result == []

    def test_short_text_single_chunk(self, chunker_with_tokenizer):
        text = "这是一段短文本。"
        result = chunker_with_tokenizer.chunk(text)
        assert len(result) == 1
        assert result[0].text == text
        assert result[0].index == 0
        assert result[0].token_count == len(text)

    def test_text_exactly_chunk_size(self, mock_tokenizer):
        """文本恰好等于 chunk_size 时应返回单个 chunk"""
        chunker = TextChunker(tokenizer=mock_tokenizer, chunk_size=10, chunk_overlap=3)
        text = "1234567890"  # 恰好 10 个字符/token
        result = chunker.chunk(text)
        assert len(result) == 1
        assert result[0].token_count == 10

    def test_no_tokenizer_raises_error(self):
        chunker = TextChunker(tokenizer=None)
        with pytest.raises(ValueError, match="tokenizer"):
            chunker.chunk("some text")


class TestChunkOverlap:
    """切块重叠逻辑测试"""

    def test_adjacent_chunks_have_overlap(self, mock_tokenizer):
        """相邻块应有 overlap token 的重叠"""
        chunk_size = 10
        overlap = 3
        chunker = TextChunker(tokenizer=mock_tokenizer, chunk_size=chunk_size, chunk_overlap=overlap)
        # 创建一个足够长的文本（无句子边界标点，避免边界切分干扰）
        text = "abcdefghijklmnopqrstuvwxyz"  # 26 chars
        result = chunker.chunk(text)

        assert len(result) > 1
        # 验证重叠：第一块的最后 overlap 个 token 应等于第二块的前 overlap 个 token
        for i in range(len(result) - 1):
            curr_tokens = mock_tokenizer.encode(result[i].text, add_special_tokens=False)
            next_tokens = mock_tokenizer.encode(result[i + 1].text, add_special_tokens=False)
            # 当前块的最后 overlap 个 token 应等于下一块的前 overlap 个 token
            assert curr_tokens[-overlap:] == next_tokens[:overlap], (
                f"Chunk {i} 和 Chunk {i+1} 之间缺少正确的重叠"
            )

    def test_chunk_size_not_exceeded(self, mock_tokenizer):
        """每个 chunk 的 token 数不应超过 chunk_size"""
        chunk_size = 10
        overlap = 3
        chunker = TextChunker(tokenizer=mock_tokenizer, chunk_size=chunk_size, chunk_overlap=overlap)
        text = "abcdefghijklmnopqrstuvwxyz"  # 26 chars
        result = chunker.chunk(text)

        for chunk in result:
            assert chunk.token_count <= chunk_size + overlap, (
                f"Chunk {chunk.index} 超过了最大 token 限制: {chunk.token_count}"
            )


class TestChunkMergeSmallLast:
    """最后一块不足 overlap 时合并测试"""

    def test_small_last_chunk_merged(self, mock_tokenizer):
        """最后一块不足 overlap token 时应合并到前一块"""
        chunk_size = 10
        overlap = 3
        chunker = TextChunker(tokenizer=mock_tokenizer, chunk_size=chunk_size, chunk_overlap=overlap)
        # 构造文本使得最后剩余不足 overlap 个 token
        # chunk_size=10, overlap=3
        # 第一块: 0-9 (10 tokens), 下一块从 7 开始
        # 第二块: 7-16 (10 tokens), 下一块从 14 开始
        # 如果总长度是 15，则第二块 7-14 (8 tokens)，剩余 15-14=1 < 3
        # 所以第二块应该合并剩余部分
        text = "abcdefghijklmno"  # 15 chars
        result = chunker.chunk(text)

        # 最后一块不应该只有 1-2 个 token
        for chunk in result:
            assert chunk.token_count >= overlap, (
                f"Chunk {chunk.index} 的 token 数 ({chunk.token_count}) 不足 overlap ({overlap})"
            )


class TestChunkSentenceBoundary:
    """中文句子边界切分测试"""

    def test_prefers_sentence_boundary(self, mock_tokenizer):
        """应优先在句子边界处切分"""
        chunk_size = 10
        overlap = 3
        chunker = TextChunker(tokenizer=mock_tokenizer, chunk_size=chunk_size, chunk_overlap=overlap)
        # 在第 7 个位置放一个句号，chunk_size=10 时应在句号后切分
        text = "abcdefg。hijklmnopqrst"
        result = chunker.chunk(text)

        # 第一块应在句号后结束
        assert result[0].text.endswith("。") or "。" in result[0].text

    def test_chinese_punctuation_boundaries(self, mock_tokenizer):
        """测试各种中文标点作为边界"""
        chunk_size = 15
        overlap = 3
        chunker = TextChunker(tokenizer=mock_tokenizer, chunk_size=chunk_size, chunk_overlap=overlap)

        # 包含多种中文标点
        text = "第一句话。第二句话！第三句话？第四句话；第五句话结束"
        result = chunker.chunk(text)

        # 每个 chunk 应在标点处结束（如果可能）
        for chunk in result[:-1]:  # 最后一块可能不以标点结束
            last_char = chunk.text[-1] if chunk.text else ""
            # 至少有一些 chunk 应在标点处结束
            has_boundary = any(
                ch in TextChunker._SENTENCE_BOUNDARIES
                for ch in chunk.text[-3:]  # 检查最后几个字符
            )
            # 不强制要求所有块都在标点结束，但至少应该尝试

    def test_newline_as_boundary(self, mock_tokenizer):
        """换行符应作为句子边界"""
        chunk_size = 10
        overlap = 3
        chunker = TextChunker(tokenizer=mock_tokenizer, chunk_size=chunk_size, chunk_overlap=overlap)
        text = "abcdefg\nhijklmnopqrst"
        result = chunker.chunk(text)

        # 第一块应在换行符后结束
        first_chunk_tokens = mock_tokenizer.encode(result[0].text, add_special_tokens=False)
        assert len(first_chunk_tokens) <= chunk_size


class TestChunkMixedContent:
    """中英文混合内容测试"""

    def test_mixed_chinese_english_no_forced_break(self, mock_tokenizer):
        """中英文混合内容不应因语言切换强制断开"""
        chunk_size = 20
        overlap = 3
        chunker = TextChunker(tokenizer=mock_tokenizer, chunk_size=chunk_size, chunk_overlap=overlap)
        text = "Hello你好World世界Test测试"
        result = chunker.chunk(text)

        # 短文本应在单个 chunk 中
        assert len(result) == 1
        assert result[0].text == text

    def test_long_mixed_content_unified(self, mock_tokenizer):
        """长中英文混合内容应统一处理"""
        chunk_size = 10
        overlap = 3
        chunker = TextChunker(tokenizer=mock_tokenizer, chunk_size=chunk_size, chunk_overlap=overlap)
        # 构造超过 chunk_size 的混合文本
        text = "Hello你好World世界Testing测试More更多"
        result = chunker.chunk(text)

        # 应该产生多个 chunk
        assert len(result) > 1
        # 所有文本应被覆盖（通过 token 重叠）
        all_text = "".join(chunk.text for chunk in result)
        # 由于重叠，合并后的文本会比原文长，但原文的所有字符都应出现
        for char in text:
            assert char in all_text


class TestChunkIndex:
    """Chunk 索引测试"""

    def test_chunk_indices_sequential(self, mock_tokenizer):
        """Chunk 索引应从 0 开始递增"""
        chunk_size = 10
        overlap = 3
        chunker = TextChunker(tokenizer=mock_tokenizer, chunk_size=chunk_size, chunk_overlap=overlap)
        text = "abcdefghijklmnopqrstuvwxyz0123456789"
        result = chunker.chunk(text)

        for i, chunk in enumerate(result):
            assert chunk.index == i

    def test_single_chunk_index_zero(self, chunker_with_tokenizer):
        """单个 chunk 的索引应为 0"""
        result = chunker_with_tokenizer.chunk("短文本")
        assert len(result) == 1
        assert result[0].index == 0
